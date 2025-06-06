from typing import Iterable


def save_markdown(summary: str, path: str) -> None:
    """Save summary text to a Markdown file."""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(summary)


def save_docx(summary: str, path: str) -> bool:
    """Save summary to a DOCX file if python-docx is available."""
    try:
        from docx import Document
    except ImportError:
        return False

    doc = Document()
    for line in summary.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('#'):
            doc.add_heading(stripped.lstrip('#').strip(), level=1)
        else:
            doc.add_paragraph(stripped)
    doc.save(path)
    return True
